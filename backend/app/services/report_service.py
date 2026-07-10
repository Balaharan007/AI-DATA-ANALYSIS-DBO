from __future__ import annotations

import io
from typing import Any, Optional

import matplotlib
matplotlib.use("Agg")
from matplotlib import pyplot as plt
import numpy as np
import pandas as pd

from fastapi import HTTPException
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, PageBreak,
)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .. import groq_client
from ..config import settings
from ..storage import new_id, now_iso, store
from . import analysis, dataset_service
from .google_drive_service import google_drive_service


def _insights_text(df, dataset_name: str, prompt: Optional[str]) -> str:
    """Generate end-to-end analysis with 5 key bullet points."""
    context = analysis.df_profile_text(df)
    ask = prompt or "Analyze this dataset end-to-end and provide key insights."
    resp = groq_client.chat_completion(
        [
            {
                "role": "system",
                "content": (
                    "You are a senior data analyst writing a detailed dataset analysis report. "
                    "Organize your answer into these EXACT sections separated by double-newlines.\n"
                    "Do NOT use markdown or bullet characters. Just plain text.\n\n"
                    "EXECUTIVE SUMMARY: (2-3 sentences on what this dataset contains and macro-level takeaway)\n"
                    "KEY METRICS: (exactly 5 bullet points, each starting with '- ', covering the most important "
                    "numeric findings, trends, and outliers. Bold the metric name in each using **bold**.)\n"
                    "INSIGHTS: (describe 2-3 interesting patterns or correlations in the data)\n"
                    "ANOMALIES: (describe any outliers, missing data issues, or data quality concerns)\n"
                    "RECOMMENDATIONS: (3 actionable business recommendations)\n\n"
                    "200-400 words total. Be precise. Use actual numbers from the data."
                ),
            },
            {"role": "user", "content": f"Dataset: {dataset_name}\n{context}\n\nTask: {ask}"},
        ],
        temperature=0.4,
        max_tokens=700,
    )
    return resp


def _parse_insights(text: str) -> dict[str, str]:
    """Parse the structured insights text into sections."""
    sections = {}
    current_key = "Executive Summary"
    current_lines: list[str] = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        upper = line.upper()
        if upper.startswith("EXECUTIVE SUMMARY"):
            if current_lines:
                sections[current_key] = "\n".join(current_lines)
            current_key = "Executive Summary"
            current_lines = [line.split(":", 1)[-1].strip()] if ":" in line else []
        elif upper.startswith("KEY METRICS"):
            if current_lines:
                sections[current_key] = "\n".join(current_lines)
            current_key = "Key Metrics"
            current_lines = [line.split(":", 1)[-1].strip()] if ":" in line else []
        elif upper.startswith("INSIGHTS"):
            if current_lines:
                sections[current_key] = "\n".join(current_lines)
            current_key = "Insights"
            current_lines = [line.split(":", 1)[-1].strip()] if ":" in line else []
        elif upper.startswith("ANOMALIES"):
            if current_lines:
                sections[current_key] = "\n".join(current_lines)
            current_key = "Anomalies"
            current_lines = [line.split(":", 1)[-1].strip()] if ":" in line else []
        elif upper.startswith("RECOMMENDATIONS"):
            if current_lines:
                sections[current_key] = "\n".join(current_lines)
            current_key = "Recommendations"
            current_lines = [line.split(":", 1)[-1].strip()] if ":" in line else []
        else:
            current_lines.append(line)
    if current_lines:
        sections[current_key] = "\n".join(current_lines)
    return sections


def _generate_chart_insights(df, chart_spec: dict) -> list[str]:
    """Generate 3 bullet-point insights for a chart (heuristic - no LLM call for speed)."""
    chart_title = chart_spec.get("title", "Untitled")
    chart_type = chart_spec.get("type", "chart")
    x_key = chart_spec.get("x_key", "N/A")
    y_keys = chart_spec.get("y_keys", [])

    # Return heuristic insights instead of LLM call
    return [
        f"This {chart_type} chart shows the distribution of {y_keys[0] if y_keys else 'values'} across different {x_key} categories.",
        f"The chart highlights which {x_key} categories have the highest and lowest {y_keys[0] if y_keys else 'values'}, making comparisons easy.",
        f"Use this visualization to identify trends, outliers, and patterns in the {x_key} dimension."
    ][:3]


def _render_chart_image(chart_spec: dict) -> io.BytesIO:
    """Render a chart spec into a PNG image in memory using matplotlib.
    Returns a BytesIO buffer containing the PNG image data."""
    plt.close("all")
    chart_type = chart_spec.get("type", "bar")
    title = chart_spec.get("title", "Chart")
    x_key = chart_spec.get("x_key", "")
    y_keys = chart_spec.get("y_keys", [])
    data = chart_spec.get("data", [])

    # Convert data to DataFrame for easier handling
    df = pd.DataFrame(data) if data else pd.DataFrame()

    # Color palette matching the frontend
    palette = ["#6366F1", "#22C55E", "#F59E0B", "#EC4899", "#06B6D4", "#EF4444", "#A855F7", "#84CC16"]

    fig, ax = plt.subplots(figsize=(6.2, 3.2))
    fig.patch.set_facecolor("#FFFFFF")
    ax.set_facecolor("#F9FAFB")

    try:
        if chart_type == "bar":
            if not df.empty and x_key in df.columns and y_keys:
                y_col = y_keys[0]
                plot_df = df.head(15)
                bars = ax.bar(range(len(plot_df)), pd.to_numeric(plot_df[y_col], errors="coerce").fillna(0),
                              color=palette[0], width=0.65, edgecolor="white", linewidth=0.5)
                ax.set_xticks(range(len(plot_df)))
                ax.set_xticklabels(plot_df[x_key].astype(str), rotation=35, ha="right", fontsize=7)
                ax.set_ylabel(str(y_col), fontsize=8)
            else:
                ax.text(0.5, 0.5, "Insufficient data", ha="center", va="center", fontsize=12, color="gray")

        elif chart_type == "pie":
            if not df.empty and x_key in df.columns and y_keys:
                y_col = y_keys[0]
                plot_df = df.head(10)
                sizes = pd.to_numeric(plot_df[y_col], errors="coerce").fillna(0)
                labels = plot_df[x_key].astype(str)
                wedges, texts, autotexts = ax.pie(
                    sizes, labels=None, autopct="%1.1f%%", startangle=90,
                    colors=palette[:len(sizes)], pctdistance=0.75,
                    wedgeprops={"edgecolor": "white", "linewidth": 0.8},
                )
                for t in autotexts:
                    t.set_fontsize(7)
                ax.legend(labels, loc="upper left", bbox_to_anchor=(1, 1), fontsize=6.5, frameon=False)
            else:
                ax.text(0.5, 0.5, "Insufficient data", ha="center", va="center", fontsize=12, color="gray")

        elif chart_type == "line":
            if not df.empty and x_key in df.columns and y_keys:
                y_col = y_keys[0]
                plot_df = df.head(100)
                xs = range(len(plot_df))
                ys = pd.to_numeric(plot_df[y_col], errors="coerce").fillna(0)
                ax.plot(xs, ys, color=palette[0], linewidth=2, marker="o", markersize=3, markerfacecolor=palette[1])
                ax.fill_between(xs, ys, alpha=0.08, color=palette[0])
                n_labels = min(10, len(plot_df))
                step = max(1, len(plot_df) // n_labels)
                tick_positions = list(range(0, len(plot_df), step))
                ax.set_xticks(tick_positions)
                ax.set_xticklabels([str(plot_df[x_key].iloc[i])[:12] for i in tick_positions], rotation=30, ha="right", fontsize=7)
                ax.set_ylabel(str(y_col), fontsize=8)
            else:
                ax.text(0.5, 0.5, "Insufficient data", ha="center", va="center", fontsize=12, color="gray")

        elif chart_type == "histogram":
            if y_keys:
                y_col = y_keys[0]
                col_data = pd.to_numeric(pd.Series(data).apply(lambda r: r.get(y_col, 0) if isinstance(r, dict) else 0)
                                         if data else pd.Series([]), errors="coerce").dropna()
                if len(col_data) > 1:
                    ax.hist(col_data, bins=min(20, max(5, int(col_data.nunique()))),
                            color=palette[0], edgecolor="white", linewidth=0.6, alpha=0.85)
                    ax.set_xlabel(str(y_col), fontsize=8)
                    ax.set_ylabel("Frequency", fontsize=8)
                else:
                    ax.text(0.5, 0.5, "Not enough data for histogram", ha="center", va="center", fontsize=10, color="gray")
            else:
                ax.text(0.5, 0.5, "No data", ha="center", va="center", fontsize=12, color="gray")

        elif chart_type == "scatter":
            if not df.empty and x_key in df.columns and len(y_keys) >= 2:
                y1, y2 = y_keys[0], y_keys[1]
                xs = pd.to_numeric(df[y1], errors="coerce").fillna(0)
                ys = pd.to_numeric(df[y2], errors="coerce").fillna(0)
                ax.scatter(xs, ys, c=palette[0], alpha=0.5, s=12, edgecolors=palette[1], linewidth=0.3)
                ax.set_xlabel(str(y1), fontsize=8)
                ax.set_ylabel(str(y2), fontsize=8)
            elif not df.empty and x_key in df.columns and y_keys:
                y_col = y_keys[0]
                xs = range(len(df))
                ys = pd.to_numeric(df[y_col], errors="coerce").fillna(0)
                ax.scatter(xs, ys, c=palette[0], alpha=0.5, s=12)
                ax.set_xlabel(str(x_key), fontsize=8)
                ax.set_ylabel(str(y_col), fontsize=8)
            else:
                ax.text(0.5, 0.5, "Insufficient data for scatter plot", ha="center", va="center", fontsize=10, color="gray")

        else:
            # Fallback: bar chart
            if not df.empty and y_keys:
                y_col = y_keys[0]
                vals = pd.to_numeric(df[y_col], errors="coerce").fillna(0)
                ax.bar(range(len(vals)), vals, color=palette[0])
                ax.set_ylabel(str(y_col), fontsize=8)
            else:
                ax.text(0.5, 0.5, "No data", ha="center", va="center", fontsize=12, color="gray")

    except Exception:
        ax.text(0.5, 0.5, "Error rendering chart", ha="center", va="center", fontsize=12, color="gray")

    ax.set_title(title, fontsize=10, fontweight="bold", pad=8, color="#1f2937")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.tick_params(axis="both", labelsize=7)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_report(dataset_id: str, prompt: Optional[str] = None) -> dict:
    """Generate a dataset analysis report with Visual Analysis (2 charts: bar + line).
    Produces both PDF (via ReportLab) and DOCX (via python-docx)."""
    df, records = dataset_service.get_context_df(dataset_id)
    if df is None or df.empty:
        raise HTTPException(404, "Dataset not found or empty.")
    dataset_name = records[0]["name"] if records else "Dataset"

    # --- 1. End-to-end analysis via LLM ---
    insights = _insights_text(df, dataset_name, prompt)
    parsed = _parse_insights(insights)

    # --- 2. Generate 2 charts for Visual Analysis (bar + line) ---
    # Use simple heuristic feature selection (no LLM calls for speed)
    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    cat_cols = [c for c in df.columns if c not in numeric_cols]

    # Filter out ID-like columns
    def is_id_like(col: str) -> bool:
        cl = col.lower()
        id_patterns = ["id", "roll", "sno", "serial", "index", "row", "number", "no", "key", "code", "uid", "pk", "fk"]
        return any(pat in cl for pat in id_patterns)

    meaningful_numeric = [c for c in numeric_cols if not is_id_like(c)]
    meaningful_cat = [c for c in cat_cols if not is_id_like(c)]

    if not meaningful_numeric:
        meaningful_numeric = numeric_cols
    if not meaningful_cat:
        meaningful_cat = cat_cols

    charts: list[dict[str, Any]] = []

    # Chart 1: Bar chart - first meaningful numeric by first meaningful categorical
    if meaningful_cat and meaningful_numeric:
        bar_spec = analysis.build_chart_spec(
            df,
            f"Bar chart showing {meaningful_numeric[0]} by {meaningful_cat[0]}",
            "bar"
        )
        # Add simple heuristic insights (no LLM call)
        bar_spec["bullet_points"] = [
            f"This bar chart compares {meaningful_numeric[0]} across different {meaningful_cat[0]} categories.",
            f"The chart highlights which {meaningful_cat[0]} categories have the highest and lowest {meaningful_numeric[0]}.",
            f"Use this to identify top-performing segments and outliers in {meaningful_numeric[0]}."
        ]
        charts.append(bar_spec)

    # Chart 2: Line chart - second meaningful numeric by first meaningful categorical (or trend)
    if len(meaningful_numeric) > 1 and meaningful_cat:
        line_spec = analysis.build_chart_spec(
            df,
            f"Line chart showing {meaningful_numeric[1]} trend over {meaningful_cat[0]}",
            "line"
        )
        line_spec["bullet_points"] = [
            f"This line chart shows the trend of {meaningful_numeric[1]} across {meaningful_cat[0]}.",
            f"The chart reveals patterns, growth, or decline in {meaningful_numeric[1]} across categories.",
            f"Use this to spot trends and compare {meaningful_numeric[1]} performance across segments."
        ]
        charts.append(line_spec)
    elif meaningful_numeric and meaningful_cat:
        # Fallback: use first numeric for line chart
        line_spec = analysis.build_chart_spec(
            df,
            f"Line chart showing {meaningful_numeric[0]} trend over {meaningful_cat[0]}",
            "line"
        )
        line_spec["bullet_points"] = [
            f"This line chart shows the trend of {meaningful_numeric[0]} across {meaningful_cat[0]}.",
            f"The chart reveals patterns and variations in {meaningful_numeric[0]} across categories.",
            f"Use this to identify trends and compare performance across segments."
        ]
        charts.append(line_spec)

    # Fallback: if no charts generated, create a simple bar chart
    if not charts and meaningful_numeric:
        spec = analysis.build_chart_spec(df, f"Bar chart of {meaningful_numeric[0]}", "bar")
        spec["bullet_points"] = [
            f"This chart shows the distribution of {meaningful_numeric[0]}.",
            "Compare values to identify high and low performers.",
            "Use for quick visual assessment of the data."
        ]
        charts.append(spec)

    # --- 3. Build PDF and DOCX ---
    report_id = new_id("rep_")
    pdf_filename = f"report_of_{dataset_name}.pdf"
    docx_filename = f"report_of_{dataset_name}.docx"
    pdf_path = settings.reports_dir / pdf_filename
    docx_path = settings.reports_dir / docx_filename

    _build_pdf(df, parsed, dataset_name, charts, records, pdf_path)
    _build_docx(df, parsed, dataset_name, charts, docx_path)

    record = {
        "id": report_id,
        "name": f"report_of_{dataset_name}",
        "created_at": now_iso(),
        "dataset_id": dataset_id,
        "dataset_name": dataset_name,
        "url": f"/reports-files/{pdf_filename}",
        "docx_url": f"/reports-files/{docx_filename}",
        "preview_url": f"/reports-files/{pdf_filename}",
        "drive_url": None,
        "drive_pdf_id": None,
        "drive_docx_id": None,
    }

    # Upload to Google Drive if enabled (PDF only)
    if settings.google_drive_enabled:
        try:
            # Read PDF content
            pdf_content = pdf_path.read_bytes()
            pdf_result = google_drive_service.upload_pdf(pdf_content, pdf_filename)
            if pdf_result:
                record["drive_url"] = pdf_result.get("webViewLink")
                record["drive_pdf_id"] = pdf_result.get("id")
        except Exception as e:
            print(f"Failed to upload PDF to Google Drive: {e}")

    store.add_report(record)
    store.add_history({
        "id": new_id("hist_"),
        "kind": "report",
        "title": record["name"],
        "created_at": record["created_at"],
        "meta": {"report_id": report_id},
    })
    return record


def _build_pdf(df, parsed, dataset_name, charts, records, path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleX", parent=styles["Title"], fontSize=20, spaceAfter=4)
    subtitle_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#6B7280"), spaceAfter=16)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, spaceBefore=18, spaceAfter=6, textColor=colors.HexColor("#1f2937"))
    h3_style = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=12, spaceBefore=12, spaceAfter=4, textColor=colors.HexColor("#374151"))
    body_style = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=9.5, leading=14, spaceAfter=8)
    bullet_style = ParagraphStyle("Bullet", parent=body_style, leftIndent=16, bulletIndent=6, spaceBefore=2, spaceAfter=2)
    chart_bullet_style = ParagraphStyle("ChartBullet", parent=body_style, leftIndent=24, fontSize=9, leading=13, spaceBefore=1, spaceAfter=1)
    small_style = ParagraphStyle("Small", parent=styles["BodyText"], fontSize=8, textColor=colors.HexColor("#6B7280"))
    doc = SimpleDocTemplate(str(path), pagesize=letter, topMargin=0.6*inch, bottomMargin=0.6*inch, leftMargin=0.6*inch, rightMargin=0.6*inch)
    story = []
    story.append(Paragraph("Insight Engine - AI Data Analyst Report", title_style))
    story.append(Paragraph("Dataset: " + dataset_name, subtitle_style))
    story.append(Paragraph("Generated: " + now_iso().split("T")[0], small_style))
    story.append(Spacer(1, 18))
    story.append(Paragraph("1. Executive Summary", h2_style))
    exec_text = parsed.get("Executive Summary", "No summary available.")
    story.append(Paragraph(exec_text.replace("\n", "<br/>"), body_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph("2. Dataset Overview", h2_style))
    overview_data = [["Metric","Value"],["Rows",str(df.shape[0])],["Columns",str(df.shape[1])],["Missing",str(int(df.isna().sum().sum()))],["Duplicates",str(int(df.duplicated().sum()))]]
    qs = str(records[0].get("quality_score","N/A")) + "%" if records else "N/A"
    overview_data.append(["Quality Score", qs])
    t = Table(overview_data, hAlign="LEFT", colWidths=[2.5*inch, 3*inch])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1f2937")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.5,colors.grey),("FONTSIZE",(0,0),(-1,-1),9),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#F9FAFB"),colors.white]),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
    story.append(t)
    story.append(Spacer(1, 12))
    story.append(Paragraph("3. Key Metrics", h2_style))
    for point in parsed.get("Key Metrics","").split("\n"):
        p = point.strip().lstrip("- ")
        if p and len(p) > 3:
            story.append(Paragraph("* " + p, bullet_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph("4. Insights & Patterns", h2_style))
    story.append(Paragraph(parsed.get("Insights","").replace("\n","<br/>"), body_style))
    story.append(Spacer(1, 8))

    # --- 5. Visual Analysis with 2 charts (bar, line) ---
    story.append(PageBreak())
    story.append(Paragraph("5. Visual Analysis", h2_style))
    story.append(Paragraph("The following 2 charts compare key features in the dataset, with 3 key insights each.", body_style))
    story.append(Spacer(1, 6))

    for i, c in enumerate(charts, 1):
        chart_title = c.get("title", "Untitled")
        chart_type = c.get("type", "chart")
        x_key = c.get("x_key", "N/A")
        y_keys = c.get("y_keys", [])

        # Chart header
        story.append(Paragraph("Chart {}: {}".format(i, chart_title), h3_style))
        story.append(Paragraph("Type: {} | X: {} | Y: {}".format(chart_type, x_key, ", ".join(str(k) for k in y_keys)), small_style))
        story.append(Spacer(1, 4))

        # Render chart image from spec
        try:
            img_buf = _render_chart_image(c)
            img = RLImage(img_buf, width=5.8*inch, height=2.8*inch)
            story.append(img)
        except Exception:
            story.append(Paragraph("[Chart rendering unavailable]", body_style))

        story.append(Spacer(1, 4))

        # 3 bullet points
        bullets = c.get("bullet_points", [])
        if not bullets:
            bullets = [
                f"This {chart_type} chart shows distribution of {y_keys[0] if y_keys else 'values'} by {x_key}.",
                "Compare categories to identify which segments perform best.",
                "Use for data-driven decision making across shown dimensions.",
            ]
        for bp in bullets:
            story.append(Paragraph("* " + bp, chart_bullet_style))

        story.append(Spacer(1, 10))
        # Page break after each chart except the last
        if i < len(charts):
            story.append(PageBreak())

    # --- 6. Recommendations (was 7) ---
    story.append(Paragraph("6. Recommendations", h2_style))
    for point in parsed.get("Recommendations","").split("\n"):
        p = point.strip().lstrip("- ")
        if p and len(p) > 3:
            story.append(Paragraph("* " + p, bullet_style))
    doc.build(story)


def _build_docx(df, parsed, dataset_name, charts, path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"
    for level in range(1, 4):
        hs = doc.styles["Heading {}".format(level)]
        hs.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    doc.add_heading("Insight Engine - AI Data Analyst Report", level=0)
    doc.add_paragraph("Dataset: " + dataset_name)
    doc.add_paragraph("Generated: " + now_iso().split("T")[0])
    doc.add_paragraph()
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(parsed.get("Executive Summary", "No summary available."))
    doc.add_heading("2. Dataset Overview", level=1)
    table = doc.add_table(rows=6, cols=2, style="Light List Accent 1")
    data = [("Rows", str(df.shape[0])), ("Columns", str(df.shape[1])), ("Missing Values", str(int(df.isna().sum().sum()))), ("Duplicates", str(int(df.duplicated().sum()))), ("Quality Score", "N/A")]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_heading("3. Key Metrics", level=1)
    for point in parsed.get("Key Metrics","").split("\n"):
        p = point.strip().lstrip("- ")
        if p and len(p) > 3:
            doc.add_paragraph(p, style="List Bullet")
    doc.add_heading("4. Insights & Patterns", level=1)
    doc.add_paragraph(parsed.get("Insights",""))

    # --- 5. Visual Analysis with 2 charts (bar, line) ---
    doc.add_heading("5. Visual Analysis", level=1)
    doc.add_paragraph("The following 2 charts were generated to compare key features in the dataset. Each chart includes 3 key analytical observations.")
    for i, c in enumerate(charts, 1):
        chart_title = c.get("title", "Untitled")
        chart_type = c.get("type", "chart")
        x_key = c.get("x_key", "N/A")
        y_keys = c.get("y_keys", [])
        doc.add_heading("Chart {}: {}".format(i, chart_title), level=2)
        doc.add_paragraph("Type: {} | X-axis: {} | Y-axis: {}".format(chart_type, x_key, ", ".join(str(k) for k in y_keys)))
        bullets = c.get("bullet_points", [])
        if not bullets:
            bullets = [
                f"This {chart_type} chart shows the distribution of {y_keys[0] if y_keys else 'values'} by {x_key}.",
                "Compare the categories to identify which segments perform best.",
                "Use this for data-driven decision making across the dimensions shown.",
            ]
        for bp in bullets:
            doc.add_paragraph(bp, style="List Bullet")

    # --- 6. Recommendations (was 7) ---
    doc.add_heading("6. Recommendations", level=1)
    for point in parsed.get("Recommendations","").split("\n"):
        p = point.strip().lstrip("- ")
        if p and len(p) > 3:
            doc.add_paragraph(p, style="List Bullet")
    doc.save(str(path))